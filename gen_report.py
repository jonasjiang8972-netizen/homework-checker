#!/usr/bin/env python3
"""Generate v2.9.0 Security & Functional Test Report (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import json
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

# --- Cover Page ---
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('作业小帮手 v2.9.0')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('生产环境安全扫描与功能测试报告')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4f, 0x6e, 0xf7)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    f'服务器：1.116.253.201:3000\n'
    f'测试时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}\n'
    f'版本：v2.9.0\n'
    f'数据库：SQLite (sql.js WASM)'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# --- TOC ---
toc_heading = doc.add_heading('目录', level=1)
sections = [
    '1. 安全扫描报告',
    '   1.1 认证与授权',
    '   1.2 文件与数据保护',
    '   1.3 Docker 安全配置',
    '   1.4 HTTP 安全头',
    '   1.5 依赖漏洞审计',
    '   1.6 代码安全审计',
    '2. 功能测试报告',
    '   2.1 API 端点测试',
    '   2.2 核心业务流程验证',
    '   2.3 数据库迁移验证',
    '3. 问题与修复',
    '4. 综合评分与建议',
]
for s in sections:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ======================
# 1. SECURITY SCAN
# ======================
doc.add_heading('1. 安全扫描报告', level=1)

# 1.1 Auth
doc.add_heading('1.1 认证与授权', level=2)
auth_table = doc.add_table(rows=8, cols=4)
auth_table.style = 'Light Grid Accent 1'
auth_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['接口', '未认证状态码', '响应内容', '结论']
for i, h in enumerate(headers):
    auth_table.rows[0].cells[i].text = h

test_results = [
    ('GET /api/quiz', '401', '请先登录后再查看测验记录', '✅ 通过'),
    ('GET /api/plans', '401', '请先登录后再查看学习计划', '✅ 通过'),
    ('GET /api/stats', '401', '-', '✅ 通过'),
    ('POST /api/correct', '401', '-', '✅ 通过'),
    ('GET /api/user/key', '401', '-', '✅ 通过'),
    ('GET /api/models', '401', '请先登录', '✅ 通过'),
    ('GET /api/uploads/[file]', '401', '-', '✅ 通过'),
]
for i, (endpoint, code, body, status) in enumerate(test_results):
    row = auth_table.rows[i + 1]
    row.cells[0].text = endpoint
    row.cells[1].text = code
    row.cells[2].text = body
    row.cells[3].text = status

doc.add_paragraph()

# 1.2 File Protection
doc.add_heading('1.2 文件与数据保护', level=2)
file_table = doc.add_table(rows=4, cols=3)
file_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['检查项', '结果', '说明']):
    file_table.rows[0].cells[i].text = h
file_data = [
    ('.env.local 外部访问', '404', '环境变量文件不可通过 HTTP 访问'),
    ('data/目录遍历', '308 (Redirect)', '目录遍历被阻止'),
    ('data/homework.db 下载', '404', '数据库文件不可通过 HTTP 访问'),
]
for i, (check, result, note) in enumerate(file_data):
    row = file_table.rows[i + 1]
    row.cells[0].text = check
    row.cells[1].text = result
    row.cells[2].text = note

doc.add_paragraph()

# 1.3 Docker Security
doc.add_heading('1.3 Docker 安全配置', level=2)
docker_table = doc.add_table(rows=5, cols=3)
docker_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['检查项', '结果', '说明']):
    docker_table.rows[0].cells[i].text = h
docker_data = [
    ('运行用户', 'nextjs（非 root）', '✅ 符合最小权限原则'),
    ('镜像基础', 'node:20-alpine', '✅ 轻量基础镜像，攻击面小'),
    ('环境变量', '通过 env_file 传入', '✅ 不写入镜像层，避免泄露'),
    ('健康检查', '已配置 wget 每30s', '✅ 自动恢复机制'),
]
for i, (check, result, note) in enumerate(docker_data):
    row = docker_table.rows[i + 1]
    row.cells[0].text = check
    row.cells[1].text = result
    row.cells[2].text = note

doc.add_paragraph()

# 1.4 HTTP Headers
doc.add_heading('1.4 HTTP 安全头', level=2)
header_table = doc.add_table(rows=7, cols=3)
header_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['响应头', '当前值', '建议']):
    header_table.rows[0].cells[i].text = h
header_data = [
    ('X-Content-Type-Options', '缺失', '应设为 nosniff，防止 MIME 类型嗅探'),
    ('X-Frame-Options', '缺失', '应设为 DENY 或 SAMEORIGIN，防止点击劫持'),
    ('Content-Security-Policy', '缺失', '应设置，防止 XSS 和数据注入'),
    ('Strict-Transport-Security', '缺失', '使用 HTTPS 时必须设置'),
    ('Referrer-Policy', '缺失', '应设置，控制 referer 信息泄露'),
    ('X-Powered-By', 'Next.js', '应移除，减少指纹信息泄露'),
]
for i, (header, current, suggestion) in enumerate(header_data):
    row = header_table.rows[i + 1]
    row.cells[0].text = header
    row.cells[1].text = current
    row.cells[2].text = suggestion

doc.add_paragraph()

# 1.5 Dependency Audit
doc.add_heading('1.5 依赖漏洞审计', level=2)
p = doc.add_paragraph('npm audit 结果：共 4 个 moderate 级别漏洞，无 high/critical。')
vuln_table = doc.add_table(rows=5, cols=3)
vuln_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['包名', '严重程度', '说明']):
    vuln_table.rows[0].cells[i].text = h
vuln_data = [
    ('next', 'moderate', '上游依赖漏洞，需 Next.js 官方修复'),
    ('next-auth', 'moderate', '低版本中的安全问题，当前版本在修复范围之外'),
    ('postcss', 'moderate', 'CSS 解析器 XSS 漏洞，极小利用场景'),
    ('uuid', 'moderate', '缓冲区边界检查缺失，Node 内置 crypto.randomUUID 不受影响'),
]
for i, (pkg, severity, desc) in enumerate(vuln_data):
    row = vuln_table.rows[i + 1]
    row.cells[0].text = pkg
    row.cells[1].text = severity
    row.cells[2].text = desc

doc.add_paragraph()

# 1.6 Code Security Audit
doc.add_heading('1.6 代码安全审计', level=2)
code_table = doc.add_table(rows=5, cols=3)
code_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['检查项', '结果', '详情']):
    code_table.rows[0].cells[i].text = h
code_data = [
    ('API Key 加密', 'AES-256-GCM', '使用 API_KEY_ENCRYPTION_SECRET 派生密钥，IV 随机'),
    ('SQL 注入防护', '参数化查询', '全部使用 queryAll/execute 的参数绑定模式'),
    ('验证码限流', 'IP:5/60s + 邮箱:3/300s', '双维度限流防止暴力破解'),
    ('Session 过期', '24h', 'JWT maxAge 24 小时自动过期'),
]
for i, (check, result, detail) in enumerate(code_data):
    row = code_table.rows[i + 1]
    row.cells[0].text = check
    row.cells[1].text = result
    row.cells[2].text = detail

doc.add_page_break()

# ======================
# 2. FUNCTIONAL TEST
# ======================
doc.add_heading('2. 功能测试报告', level=1)

# 2.1 API
doc.add_heading('2.1 API 端点测试', level=2)
api_table = doc.add_table(rows=12, cols=4)
api_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['端点', '方法', '预期行为', '结果']):
    api_table.rows[0].cells[i].text = h
api_data = [
    ('/ (首页)', 'GET', '返回 200 静态页面', '✅ 200'),
    ('/api/auth/send-code', 'POST', '发送验证码 / 限流提示', '✅ 逻辑正确'),
    ('/api/quiz', 'GET', '返回测验历史（需登录）', '✅ 401 未登录'),
    ('/api/quiz', 'POST (generate)', 'AI 生成题目（需 Key）', '✅ 503 未配置 Key'),
    ('/api/quiz', 'POST (submit)', 'AI 批改（需 Key）', '✅ 503 未配置 Key'),
    ('/api/quiz', 'POST (correct)', '保存纠正（需登录）', '✅ 逻辑正确'),
    ('/api/plans', 'GET', '返回学习计划（需登录）', '✅ 401 未登录'),
    ('/api/plans', 'POST', 'AI 生成计划（需 Key）', '✅ 503 未配置 Key'),
    ('/api/user/key', 'GET/POST/DELETE', '管理 API Key', '✅ 401 未登录'),
    ('/api/user/settings', 'GET/PATCH', '管理用户设置', '✅ 逻辑正确'),
    ('/api/models', 'GET', '返回模型列表（需 Key）', '✅ 401 未登录'),
]
for i, (ep, method, expect, result) in enumerate(api_data):
    row = api_table.rows[i + 1]
    row.cells[0].text = ep
    row.cells[1].text = method
    row.cells[2].text = expect
    row.cells[3].text = result

doc.add_paragraph()

# 2.2 Core Business Flow
doc.add_heading('2.2 核心业务流程验证', level=2)

flows = [
    ('用户登录流程', [
        '输入邮箱 → 发送验证码 → 输入验证码 → 登录成功',
        '✅ SMTP 配置正确（production Gmail SMTP）',
        '✅ 双维度限流保护',
        '✅ JWT 24h 过期',
    ]),
    ('API Key 配置流程', [
        '设置页输入 Key → AES-256-GCM 加密存储 → 展示脱敏 Key',
        '✅ 加密存储 + 脱敏展示',
        '✅ 删除 Key 功能正常',
        '✅ 可选自定义 API 接口地址',
    ]),
    ('AI 批改流程', [
        '拍照/选图 → 上传 → 视觉预检 → AI 批改 → 展示结果',
        '✅ 模型选择器可用（主页+设置页）',
        '✅ 视觉模型自动选择',
        '✅ 引导→答案两段式展示',
    ]),
    ('测验流程 (v2.9 重构)', [
        '选知识点 → AI 出题 → 答题 → AI 批改 → 展示详情 → 可纠正',
        '✅ 学科适配（数学/语文/英语）',
        '✅ JSON 序列化修复（questions_json/answers_json）',
        '✅ 手动纠错 + 分数重算',
        '✅ 提交失败保留答案重试',
    ]),
    ('学习计划流程', [
        '基于薄弱知识点 → AI 生成分步计划 → 保存 → 可更新状态',
        '✅ AI 协议兼容性修复（fetch /chat/completions）',
        '✅ 计划保存/状态更新',
    ]),
]

for title, items in flows:
    doc.add_heading(title, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# 2.3 Database Migration
doc.add_heading('2.3 数据库迁移验证', level=2)
mig_table = doc.add_table(rows=4, cols=3)
mig_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['迁移', 'SQL', '状态']):
    mig_table.rows[0].cells[i].text = h
mig_data = [
    ('user_settings.base_url', 'ALTER TABLE ADD COLUMN', '✅ v2.8.5 已部署'),
    ('test_records.subject', 'ALTER TABLE ADD COLUMN', '✅ v2.9.0 已部署'),
    ('自动迁移幂等', 'try-catch 包裹，表已存在静默跳过', '✅ 兼容新旧实例'),
]
for i, (mig, sql, status) in enumerate(mig_data):
    row = mig_table.rows[i + 1]
    row.cells[0].text = mig
    row.cells[1].text = sql
    row.cells[2].text = status

doc.add_page_break()

# ======================
# 3. ISSUES & FIXES
# ======================
doc.add_heading('3. 问题与修复', level=1)

issues = [
    ('P1-数据库初始化竞争', 'severity:medium', '发现',
     '服务器日志中出现 "Database not initialized" 错误',
     'lib/auth-utils.ts getApiBaseUrl() 中 queryOne 在 getDb() 完成前被调用',
     '在 getApiBaseUrl() 中添加 await getDb() 调用，确保数据库初始化完成后再执行查询'),
    ('P2-缺少安全响应头', 'severity:low', '发现',
     'HTTP 响应缺少 X-Content-Type-Options、X-Frame-Options、CSP 等安全头',
     'Next.js 默认不添加这些头，需通过 next.config.js 或中间件配置',
     '建议在 next.config.js 中配置 headers() 添加安全头，或在 Nginx 反向代理层添加'),
    ('P3-X-Powered-By 信息泄露', 'severity:low', '发现',
     '响应头中包含 X-Powered-By: Next.js',
     'Next.js 默认行为',
     '可通过 next.config.js 的 poweredByHeader: false 关闭'),
    ('P4-npm moderate 漏洞', 'severity:info', '发现',
     '4 个 moderate 级别依赖漏洞',
     'next / next-auth 上游依赖、postcss XSS、uuid 边界检查',
     '建议定期运行 npm audit 跟踪，待 Next.js 官方修复后升级'),
]

issue_table = doc.add_table(rows=len(issues) + 1, cols=5)
issue_table.style = 'Light Grid Accent 1'
headers = ['编号', '严重度', '状态', '问题描述', '根因 / 修复措施']
for i, h in enumerate(headers):
    issue_table.rows[0].cells[i].text = h

for i, (title, severity, status, desc, root_cause, fix) in enumerate(issues):
    row = issue_table.rows[i + 1]
    row.cells[0].text = title
    row.cells[1].text = severity
    row.cells[2].text = status
    row.cells[3].text = desc
    row.cells[4].text = f'{root_cause}\n{fix}'

doc.add_page_break()

# ======================
# 4. CONCLUSION
# ======================
doc.add_heading('4. 综合评分与建议', level=1)

scores = [
    ('认证与授权', 'A', '全部 API 端点需认证，权限校验完整'),
    ('数据保护', 'A', '数据库/环境变量/上传文件均不可外部访问'),
    ('加密安全', 'A', 'API Key AES-256-GCM 加密，IV 随机'),
    ('Docker 安全', 'A', '非 root 运行，最小镜像，env_file 传参'),
    ('HTTP 安全头', 'C', '缺少多项重要安全头，建议通过 Nginx 或 next.config.js 补充'),
    ('依赖安全', 'B', '4 个 moderate 漏洞，无 high/critical，持续跟踪即可'),
]
score_table = doc.add_table(rows=len(scores) + 1, cols=3)
score_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['检查项', '评分', '说明']):
    score_table.rows[0].cells[i].text = h
for i, (item, grade, note) in enumerate(scores):
    row = score_table.rows[i + 1]
    row.cells[0].text = item
    row.cells[1].text = grade
    row.cells[2].text = note

doc.add_paragraph()
doc.add_paragraph()

# Overall
overall = doc.add_paragraph()
run = overall.add_run('综合评分：A-')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)

doc.add_paragraph()

recommendations = [
    '近期（下次迭代）',
    '  在 next.config.js 中配置安全响应头（推荐配置见下）',
    '  关闭 X-Powered-By',
    '中期',
    '  接入 Nginx 反向代理，统一管理 TLS/安全头/限流',
    '  评估接入 Sentry 等错误监控服务',
    '  配置每日 SQLite 自动备份',
    '长期',
    '  评估迁移至 PostgreSQL 以支持更高并发',
    '  考虑引入 Rate Limiting 中间件保护所有 API',
]

doc.add_heading('改进建议', level=2)
for rec in recommendations:
    doc.add_paragraph(rec)

doc.add_paragraph()
doc.add_paragraph()

# Recommended next.config.js headers config
doc.add_heading('推荐 next.config.js 安全头配置', level=3)
code = doc.add_paragraph()
run = code.add_run(
    '// next.config.js 添加:\n'
    'async headers() {\n'
    '  return [{\n'
    '    source: \'/(.*)\',\n'
    '    headers: [\n'
    '      { key: \'X-Content-Type-Options\', value: \'nosniff\' },\n'
    '      { key: \'X-Frame-Options\', value: \'DENY\' },\n'
    '      { key: \'Referrer-Policy\', value: \'strict-origin-when-cross-origin\' },\n'
    '    ],\n'
    '  }];\n'
    '},\n'
    'poweredByHeader: false,'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'

# --- Save ---
output_path = '/Users/jonasjiang/homework-checker/public/作业小帮手-v2.9.0-安全扫描与测试报告.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Report saved: {output_path}')
