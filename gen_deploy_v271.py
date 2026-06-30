# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

FONT = 'Microsoft YaHei'; MONO = 'Courier New'

def sr(r, b=False, s=11, c=None, f=FONT, italic=False):
    r.bold = b; r.font.size = Pt(s); r.font.name = f; r.italic = italic
    r._element.rPr.rFonts.set(qn('w:eastAsia'), f)
    if c: r.font.color.rgb = RGBColor(*c)

def pa(t, b=False, s=11, c=None, a=None, sa=6, sb=0, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); sr(r, b, s, c, italic=italic)
    pf = p.paragraph_format; pf.space_after = Pt(sa); pf.space_before = Pt(sb)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.35
    if a: p.alignment = a; return p

def hd(t, l=1):
    h = doc.add_heading(t, l)
    sz = {1:20,2:15,3:12}
    for r in h.runs: sr(r, True, sz.get(l,11))
    return h

def cb(t, lb=None):
    if lb:
        p = doc.add_paragraph(); r = p.add_run(lb); sr(r, True, 10, (0x4f,0x6e,0xf7))
        p.paragraph_format.space_after = Pt(2)
    for line in t.strip().split('\n'):
        p = doc.add_paragraph(); r = p.add_run(line or ' '); sr(r, 9, f=MONO)
        pf = p.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p._element.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'))

def nt(t, tp='info'):
    cl = {'ok':(0x27,0xae,0x60),'warn':(0xe6,0x7e,0x22),'info':(0x4f,0x6e,0xf7)}
    p = doc.add_paragraph(); r = p.add_run(f'  {t}'); sr(r, s=10, c=cl.get(tp,(0x4f,0x6e,0xf7)), italic=True)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)

def sep():
    p = doc.add_paragraph(); r = p.add_run('─'*70); sr(r, s=6, c=(0xDD,0xDD,0xDD))

def tbl(hdrs, rows):
    t = doc.add_table(rows=1, cols=len(hdrs)); t.style = 'Light Grid Accent 1'
    for i,h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: sr(r, True, 9)
    for rd in rows:
        c = t.add_row().cells
        for i,v in enumerate(rd):
            c[i].text = str(v)
            for p in c[i].paragraphs:
                for r in p.runs: sr(r, 9)
    return t

# ===== COVER =====
for _ in range(3): doc.add_paragraph()
pa('📚 作业小帮手', b=1, s=30, c=(0x1a,0x1a,0x2e), a=WD_ALIGN_PARAGRAPH.CENTER)
pa('Homework Checker', s=16, c=(0x8e,0x95,0xa2), italic=1, a=WD_ALIGN_PARAGRAPH.CENTER)
pa('部署运维操作手册 v2.7.1', b=1, s=22, c=(0x4f,0x6e,0xf7), a=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3): doc.add_paragraph()

for l,v in [('版本','v2.8.0'),('更新日期','2026-06-30'),('服务器','<your-server-ip>'),('系统','Ubuntu 24.04 LTS | Docker + Nginx')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{l}：'); sr(r, 1, 11, (0x55,0x55,0x55))
    r = p.add_run(v); sr(r, s=11)

doc.add_page_break()

# ===== TOC =====
hd('目录',1)
for x in ['第一章 项目与架构','第二章 环境准备','第三章 获取代码','第四章 环境变量','第五章 Docker 部署',
          '第六章 Nginx 配置','第七章 应用更新流程','第八章 安全运维','第九章 v2.7.1 新增内容','附录']:
    pa(x, 10, sa=2)
doc.add_page_break()

# ===== 1 =====
hd('第一章 项目与架构',1)
pa('1.1 技术架构总览',b=1,s=12,sb=8)
tbl(['组件','版本','说明'],
    [['Next.js','16.2.9','App Router 全栈框架'],
     ['React','19.2.7','UI 框架'],
     ['TypeScript','6.0.3','类型安全'],
     ['SQLite','sql.js WASM','嵌入数据库'],
     ['Nodemailer','9.0.1','SMTP 邮件发送'],
     ['NextAuth','4.24.14','JWT 认证'],
     ['Vitest','4.1.9','测试框架'],
     ['Docker','29.6.1','容器化部署'],
     ['Nginx','1.24','反向代理']])

pa('1.2 部署架构',b=1,s=12,sb=8)
cb("""用户浏览器 → HTTP :80
              │
         ┌────▼────┐
         │  Nginx  │  ← 反向代理
         └────┬────┘
              │ proxy_pass 127.0.0.1:3000
         ┌────▼────┐
         │ Docker  │  ← Node.js 20 Alpine
         │ :3000   │     homework-checker-app
         └────┬────┘
              │
    ┌─────────┴─────────┐
    │  data/homework.db  │ ← SQLite 持久化
    │  data/uploads/     │ ← 🔒 API 鉴权访问
    └───────────────────┘""")

nt('v2.7.1 变更：Supabase 适配器已移除，邮件改用 nodemailer v9 驱动 Gmail SMTP。图片上传经 /api/uploads/[file] 鉴权保护。','ok')
doc.add_page_break()

# ===== 2 =====
hd('第二章 环境准备',1)
pa('2.1 SSH 连接',b=1,s=12,sb=6)
cb('ssh ubuntu@<your-server-ip>', '命令：')

pa('2.2 安装 Docker',b=1,s=12,sb=6)
cb("""sudo apt update && sudo apt upgrade -y
sudo apt install -y ca-certificates curl
sudo install -m 0755 -d /etc/apt/keyrings
sudo curl -fsSL https://download.docker.com/linux/ubuntu/gpg -o /etc/apt/keyrings/docker.asc
sudo chmod a+r /etc/apt/keyrings/docker.asc
echo "deb [arch=$(dpkg --print-architecture) signed-by=/etc/apt/keyrings/docker.asc] https://download.docker.com/linux/ubuntu $(. /etc/os-release && echo "$VERSION_CODENAME") stable" | sudo tee /etc/apt/sources.list.d/docker.list > /dev/null
sudo apt update && sudo apt install -y docker-ce docker-ce-cli containerd.io docker-buildx-plugin docker-compose-plugin""")

pa('2.3 安装 Nginx',b=1,s=12,sb=6)
cb('sudo apt install -y nginx')

pa('2.4 防火墙',b=1,s=12,sb=6)
cb("""sudo ufw allow 22/tcp
sudo ufw allow 80/tcp
sudo ufw allow 443/tcp
sudo ufw enable""")

pa('2.5 时区',b=1,s=12,sb=6)
cb('sudo timedatectl set-timezone Asia/Shanghai')

doc.add_page_break()

# ===== 3 =====
hd('第三章 获取项目代码',1)
cb("""git clone https://github.com/jonasjiang8972-netizen/homework-checker.git
cd homework-checker""", '命令：')

pa('项目结构',b=1,s=12,sb=6)
cb("""homework-checker/
├── app/api/            # API 路由（含 uploads/[file] 鉴权接口）
├── lib/                # 核心库
│   ├── email.ts        # Gmail SMTP 发送（nodemailer v9）
│   ├── encryption.ts   # AES-256-GCM 加密
│   ├── supabase.ts     # SQLite 查询构建器
│   └── db.ts           # SQLite 数据库初始化
├── data/               # 数据库 + 图片上传（Docker volume）
├── __tests__/          # 66 个测试用例
├── Dockerfile          # 多阶段构建
└── docker-compose.yml  # 编排文件""")

doc.add_page_break()

# ===== 4 =====
hd('第四章 环境变量',1)
pa('在 .env.local 中配置：',b=1,s=12,sb=6)
cb("""ANTHROPIC_API_KEY=sk-your_siliconflow_key
ANTHROPIC_BASE_URL=https://api.siliconflow.cn/v1
ANTHROPIC_MODEL=Qwen/Qwen3-VL-32B-Instruct

SMTP_HOST=smtp.gmail.com
SMTP_PORT=587
SMTP_USER=your_email@gmail.com
SMTP_PASS=your_16_char_gmail_app_password

NEXTAUTH_SECRET=$(openssl rand -hex 32)
NEXTAUTH_URL=http://<your-domain>

API_KEY_ENCRYPTION_SECRET=$(openssl rand -hex 32)""")

tbl(['变量','必填','说明','获取方式'],
    [['ANTHROPIC_API_KEY','✅','SiliconFlow 密钥','siliconflow.cn 控制台'],
     ['SMTP_HOST/USER/PASS','✅','Gmail SMTP','Google 账户→应用专用密码'],
     ['NEXTAUTH_SECRET','✅','JWT 签名密钥','openssl rand -hex 32'],
     ['API_KEY_ENCRYPTION_SECRET','✅ v2.7+','API Key 加密密钥','openssl rand -hex 32']])

doc.add_page_break()

# ===== 5 =====
hd('第五章 Docker 部署',1)
pa('5.1 Dockerfile（多阶段构建）',b=1,s=12,sb=6)
cb("""FROM node:20-alpine AS base
  ├── deps     → npm ci --only=production --legacy-peer-deps
  ├── builder  → npm ci --legacy-peer-deps + npm run build
  └── runner   → COPY --from=builder .next/standalone + node_modules""")

pa('5.2 构建与启动',b=1,s=12,sb=6)
cb("""cd /home/ubuntu/homework-checker
docker compose build
docker compose up -d
sleep 10
curl http://localhost:3000""", '命令：')

pa('5.3 验证',b=1,s=12,sb=6)
cb("""docker ps  # 确认 homework-checker-app-1 状态 healthy
curl -s -o /dev/null -w '%{http_code}' http://localhost:3000
# 返回 200""")

doc.add_page_break()

# ===== 6 =====
hd('第六章 Nginx 配置',1)
cb("""server {
    listen 80;
    server_name <your-domain>;

    proxy_connect_timeout 120s;
    proxy_send_timeout 180s;
    proxy_read_timeout 180s;

    location /reports/ {
        alias /var/www/reports/;
    }

    location / {
        proxy_pass http://127.0.0.1:3000;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection 'upgrade';
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
        proxy_cache_bypass $http_upgrade;
    }
}""", '配置文件：/etc/nginx/sites-enabled/homework-checker')

cb("""sudo nginx -t          # 测试语法
sudo systemctl reload nginx  # 重载配置""", '验证：')

doc.add_page_break()

# ===== 7 =====
hd('第七章 应用更新流程',1)
cb("""# 一行命令完成更新
cd /home/ubuntu/homework-checker \\
  && git pull \\
  && docker compose build \\
  && docker compose up -d""", '日常更新命令：')

pa('7.1 完整更新示例（安全升级）',b=1,s=12,sb=6)
cb("""# 1. 本地修改代码、测试
npm test
npm run build

# 2. 推送 GitHub
git add -A && git commit -m "本次变更说明"
git push origin main
git tag v2.7.x && git push origin v2.7.x

# 3. SSH 到服务器拉取
ssh ubuntu@<your-server-ip>
cd /home/ubuntu/homework-checker && git pull

# 4. 构建部署
docker compose build && docker compose up -d

# 5. 验证
curl http://localhost:3000
docker logs homework-checker-app-1 --tail 20""")

doc.add_page_break()

# ===== 8 =====
hd('第八章 安全运维',1)
pa('8.1 日常安全检查命令',b=1,s=12,sb=6)
cb("""# 查看容器状态
docker ps
docker inspect --format='{{.State.Health.Status}}' homework-checker-app-1

# 查看日志
docker logs homework-checker-app-1 --tail 50

# 测试认证保护
curl -s http://localhost:3000/api/uploads/test.jpg  # 应返回 401
curl -s http://localhost:3000/api/plans              # 应返回 401

# 数据库备份
cp data/homework.db data/homework.db.bak.$(date +%Y%m%d)""")

pa('8.2 v2.7.0 安全加固清单',b=1,s=12,sb=6)
tbl(['措施','文件','说明'],
    [['上传鉴权','app/api/uploads/[file]/route.ts','图片需要登录后才能访问'],
     ['加密密钥强制','lib/encryption.ts','API_KEY_ENCRYPTION_SECRET 必须配置'],
     ['验证码限流','send-code/route.ts','IP 5次/分 + 邮箱 3次/5分'],
     ['JWT 过期','[...nextauth]/route.ts','24 小时自动过期'],
     ['API 401 保护','全部路由','未登录访问返回 401']])

pa('8.3 v2.7.1 依赖升级',b=1,s=12,sb=6)
tbl(['变更','影响','漏洞变化'],
    [['移除 @auth/supabase-adapter','清除废弃的 Supabase 引用','消除 cookie/nodemailer 传递依赖漏洞'],
     ['nodemailer ^7 → ^9','SMTP 邮件驱动升级','修复 6 个高危/中危 CVE'],
     ['Dockerfile --legacy-peer-deps','兼容 npm 严格 peer 检查','保障构建通过']])

doc.add_page_break()

# ===== 9 =====
hd('第九章 v2.7.1 新增内容',1)
pa('9.1 依赖升级详情',b=1,s=12,sb=8)
tbl(['包名','v2.6.2→v2.7.0','v2.7.0→v2.7.1','说明'],
    [['nodemailer','（安装）','^7.0.13 → ^9.0.1','修复 6 个 CVE'],
     ['@auth/supabase-adapter','（保留）','已移除','不再需要 Supabase'],
     ['@supabase/supabase-js','（保留）','已移除','不再需要 Supabase']])

pa('9.2 npm audit 对比',b=1,s=12,sb=6)
tbl(['版本','高危','中危','低危','合计'],
    [['v2.6.2','1','7','1','9'],
     ['v2.7.0','1','6','1','8'],
     ['v2.7.1','0','4','0','4 ✅']])

pa('9.3 安全修复累计',b=1,s=12,sb=6)
tbl(['漏洞','版本','类型'],
    [['上传图片无认证访问','v2.7.0','安全加固'],
     ['加密密钥硬编码回退','v2.7.0','安全加固'],
     ['验证码无频率限制','v2.7.0','安全加固'],
     ['JWT 永不过期','v2.7.0','安全加固'],
     ['data/ 不在 .gitignore','v2.7.0','安全加固'],
     ['未授权访问标题/计划/统计 API','v2.6.2','安全修复'],
     ['Supabase 废弃包 + nodemailer CVE','v2.7.1','依赖升级']])

doc.add_page_break()

# ===== APPENDIX =====
hd('附录',1)
hd('A. 当前服务器状态',2)
tbl(['项目','值'],
    [['服务器','<your-server-ip>'],
     ['操作系统','Ubuntu 24.04 LTS'],
     ['Docker','29.6.1 / Compose v5.2.0'],
     ['Nginx','1.24'],
     ['Node.js（容器内）','20.20.2 Alpine'],
     ['应用版本','v2.7.1'],
     ['数据库','SQLite (sql.js) @ data/homework.db'],
     ['邮件驱动','nodemailer 9.x → Gmail SMTP'],
     ['存储路径','data/uploads/（鉴权保护）']])

hd('B. 端口说明',2)
tbl(['端口','用途','防火墙'],
    [['22','SSH','✅ 开放'],
     ['80','HTTP Nginx','✅ 开放'],
     ['443','HTTPS','✅ 开放'],
     ['3000','Docker 容器','❌ 不对外开放']])

hd('C. 相关链接',2)
tbl(['资源','地址'],
    [['GitHub','https://github.com/jonasjiang8972-netizen/homework-checker'],
     ['应用','http://<your-domain>'],
     ['测试报告','<内部文档链接>'],
     ['安全扫描报告','<内部文档链接>']])

nt('当前手册版本 v2.7.1 | 生成于 2026-06-30','info')
doc.add_paragraph()
pa('— 文档结束 · 作业小帮手 v2.7.1 部署运维操作手册 —', s=9, c=(0xAA,0xAA,0xAA), italic=1, a=WD_ALIGN_PARAGRAPH.CENTER)

out = '/Users/jonasjiang/homework-checker/public/作业小帮手v2.7.1-部署运维操作手册.docx'
doc.save(out)
print(f'Saved: {out}  Size: {os.path.getsize(out)} bytes')
